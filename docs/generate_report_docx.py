# -*- coding: utf-8 -*-
"""Generate project report as DOCX with embedded screenshots."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = Path(__file__).parent
SHOTS = ROOT / "screenshots"
OUT = ROOT / "项目报告-第8组-TalentFlow-ATS.docx"


def set_cn_font(run, name="宋体", size=12):
    run.font.name = name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_title(doc, text, size=16, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    set_cn_font(r, "黑体", size)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    set_cn_font(r, "黑体", 14)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    set_cn_font(r, "黑体", 12)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn_font(r, "宋体", 11)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_image(doc, path, caption):
    if path.exists():
        doc.add_picture(str(path), width=Inches(5.8))
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        set_cn_font(cr, "宋体", 10)
    doc.add_page_break()


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(2.5)

    add_title(doc, "智能化软件开发实践", 18)
    add_title(doc, "项目报告", 22)
    doc.add_paragraph()
    add_title(doc, "第八组", 14, False)
    doc.add_paragraph()

    members = [
        ("刘沛秋", "23009200512"),
        ("邱尹鸿", "23009201337"),
        ("冯逸凡", "23009201054"),
        ("王天睿", "23009200966"),
    ]
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "姓名"
    hdr[1].text = "学号"
    for i, (n, s) in enumerate(members, 1):
        table.rows[i].cells[0].text = n
        table.rows[i].cells[1].text = s
    doc.add_page_break()

    sections = [
        ("一、项目简介", [
            "本项目开发面向互联网行业的 TalentFlow ATS 智能招聘平台，解决 HR 筛选效率低、候选人体验差、管理决策缺乏数据支撑三大痛点。",
            "系统采用 C/B/M 三端架构：C 端服务求职者智能投递与进度追踪；B 端服务 HR 漏斗管理与 AI 匹配；M 端服务管理层数据驾驶舱。",
            "技术栈：Spring Boot 3.2 + JPA + H2/MySQL 后端，Vue 3 + Element Plus + ECharts 前端，Token 鉴权与路由守卫实现多角色隔离。",
            "核心能力包括岗位全生命周期管理、Excel 批量导入、审批工作流、AI 简历解析与匹配评分、招聘漏斗、管理驾驶舱与人才库。",
        ]),
        ("二、需求分析", [
            "系统涉及六类用户：求职者 CANDIDATE、部门账号 DEPARTMENT、招聘 HR ADMIN、面试官 INTERVIEWER、管理层 EXECUTIVE、游客。",
            "C 端用户故事：免填投递、外卖式进度追踪、24h AI 答疑、淘汰后有温度反馈与岗位推荐。",
            "B 端用户故事：AI 匹配分辅助筛选、漏斗一站式操作、结构化面评减少偏见。",
            "M 端用户故事：渠道 ROI 分析、私域人才库激活。",
            "主要用例：UC-01 登录认证、UC-02 智能投递、UC-03 漏斗推进、UC-04 岗位审批、UC-05 Excel 导入、UC-06 管理驾驶舱。",
            "测试用例 TC-001 至 TC-015 覆盖登录、投递、审批、权限隔离、AI 答疑等场景，优先级分为高/中/低三级。",
        ]),
        ("三、系统概要设计", [
            "硬件：Windows 11 笔记本，16GB 内存。软件：JDK 17、Maven 3.9+、Node.js 18+、H2/MySQL。",
            "架构：前后端分离 B/S，前端 :5180，后端 :8080，Axios 调用 /webapi REST API。",
            "鉴权链：AuthInterceptor 解析 Bearer Token → UserContextHolder → 角色校验 → Controller。",
            "数据模型：sys_user、position、application、interview、interview_evaluation 五张核心表。",
            "ER 关系：用户 1:N 申请 N:1 岗位；申请 1:N 面试与面评。",
            "功能模块：C 端公开岗位/智能投递/进度追踪；B 端岗位管理/漏斗/面评；M 端 KPI/图表/ROI。",
        ]),
    ]
    for title, paras in sections:
        add_h1(doc, title)
        for t in paras:
            add_body(doc, t)
        doc.add_page_break()

    add_h1(doc, "四、详细设计及实现")
    add_h2(doc, "4.1 后端接口规范")
    add_body(doc, "统一响应格式 { code, message, data }。认证：POST /webapi/auth/login。岗位：GET/POST/PUT/DELETE /webapi/positions/*。C 端：/webapi/candidate/apply/upload 等。B 端：/webapi/recruiter/applications。M 端：/webapi/management/dashboard。")
    add_h2(doc, "4.2 系统界面截图")
    add_body(doc, "以下截图为系统实际运行界面，由 Playwright 在 http://localhost:5180 自动截取。")

    images = [
        ("01-login.png", "图 4-1 用户登录页"),
        ("02-public-list.png", "图 4-2 公开岗位列表"),
        ("03-public-detail.png", "图 4-3 公开岗位详情"),
        ("04-candidate-applications.png", "图 4-4 我的投递（C 端进度追踪）"),
        ("05-candidate-apply.png", "图 4-5 智能投递页"),
        ("06-recruiter-pipeline.png", "图 4-6 招聘漏斗（HR）"),
        ("07-admin-positions.png", "图 4-7 岗位管理（HR）"),
        ("08-statistics.png", "图 4-8 数据统计"),
        ("09-dept-positions.png", "图 4-9 部门岗位管理"),
        ("10-excel-import.png", "图 4-10 Excel 批量导入"),
        ("11-position-create.png", "图 4-11 新建岗位表单"),
        ("12-management-dashboard.png", "图 4-12 管理驾驶舱（M 端）"),
        ("13-interviewer-pipeline.png", "图 4-13 面试官视图"),
    ]
    for fname, cap in images:
        add_image(doc, SHOTS / fname, cap)

    add_h1(doc, "4.3 核心模块实现")
    add_body(doc, "TokenStore + AuthInterceptor 实现鉴权；ResumeParseService 正则解析简历；AiMatchingService 关键词匹配计算 30-99 分；ApplicationService 管理漏斗阶段流转；authState reactive 解决登录 UI 不刷新。")
    add_h2(doc, "4.4 AI 辅助开发")
    add_body(doc, "使用 Cursor AI 编码助手完成 90% 以上代码生成，包括三端导航、CandidateApplications、RecruiterPipeline、ManagementDashboard 及后端 Service/Controller。")
    doc.add_page_break()

    add_h1(doc, "五、系统测试")
    add_h2(doc, "5.1 功能清单")
    add_body(doc, "共 20 项功能 FUN-01 至 FUN-20，涵盖登录、CRUD、审批、导入、智能投递、漏斗、面评、驾驶舱等，全部实现并通过测试。")
    add_h2(doc, "5.2 需求覆盖")
    add_body(doc, "15 项核心需求测试覆盖率 100%。界面截图测试记录见第四节，各角色页面均可正常渲染。")
    add_h2(doc, "5.3 Bug 列表")
    add_body(doc, "BUG-01 API 404（后端未重启）；BUG-02 Header 不刷新（localStorage）；BUG-03 import 路径；BUG-04/05 端口冲突；BUG-06 jar 锁定；BUG-07 bat 编码。均已解决。")
    doc.add_page_break()

    add_h1(doc, "六、总结与思考")
    add_body(doc, "本项目完成招聘岗位管理系统到 TalentFlow ATS 的演进，作业必做功能 100% 覆盖，并扩展 C/B/M 三端与 AI 辅助能力。")
    add_body(doc, "心得：三端架构分工清晰；AI 辅助编程大幅提升效率；响应式状态与路由守卫是多角色前端关键；规则引擎可快速模拟 AI 能力。")
    add_body(doc, "启动：双击 start-all.bat，访问 http://localhost:5180。测试账号 candidate/candidate123、admin/admin123、executive/exec123 等。")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("本报告由第 8 组全体成员共同完成。\n刘沛秋 · 邱尹鸿 · 冯逸凡 · 王天睿")
    set_cn_font(r, "宋体", 12)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
